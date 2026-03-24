import pandas as pd
import numpy as np
from docx import Document
from fpdf import FPDF
import zipfile, os
from datetime import date

out = 'C:/Users/kerry/repos/mgmt675/files/m4_files'
for ex in ['Example1','Example2','Example3']:
    os.makedirs(f'{out}/{ex}', exist_ok=True)

# ── Example1: portfolio.xlsx ──────────────────────────────────────────────────
holdings = [
    ('AAPL','Technology','Equities',2340000),
    ('MSFT','Technology','Equities',1980000),
    ('GOOGL','Technology','Equities',1450000),
    ('NVDA','Technology','Equities',1620000),
    ('AMZN','Consumer Cyclical','Equities',1230000),
    ('TSLA','Consumer Cyclical','Equities',780000),
    ('HD','Consumer Cyclical','Equities',540000),
    ('JPM','Financial Services','Equities',1100000),
    ('BAC','Financial Services','Equities',680000),
    ('GS','Financial Services','Equities',590000),
    ('JNJ','Healthcare','Equities',870000),
    ('UNH','Healthcare','Equities',1050000),
    ('PFE','Healthcare','Equities',420000),
    ('XOM','Energy','Equities',640000),
    ('CVX','Energy','Equities',480000),
    ('AGG','Fixed Income','Bonds',1500000),
    ('LQD','Fixed Income','Bonds',900000),
    ('TLT','Fixed Income','Bonds',600000),
    ('GLD','Commodities','Alternatives',750000),
    ('VNQ','Real Estate','Alternatives',430000),
    ('BTC-USD','Digital Assets','Alternatives',380000),
]
df = pd.DataFrame(holdings, columns=['Ticker','Sector','Asset Class','Market Value'])
df.to_excel(f'{out}/Example1/portfolio.xlsx', index=False)
print(f"Example1/portfolio.xlsx: {len(df)} rows")

# ── Example2: loan_terms.docx + payments.xlsx ─────────────────────────────────
borrowers = [
    ('Acme Industrial LLC',    'Commercial RE', '2021-03-15', 1200000, 5.25),
    ('Sunrise Hospitality',    'Commercial RE', '2020-07-01',  850000, 4.75),
    ('BlueSky Logistics',      'Commercial RE', '2022-01-10', 2100000, 5.50),
    ('Greenfield Partners',    'Construction',  '2021-11-20',  650000, 6.00),
    ('Harbor View Apartments', 'Multifamily',   '2020-04-05', 3400000, 4.50),
    ('Metro Storage Inc',      'Commercial RE', '2023-02-14',  420000, 6.25),
    ('Pinnacle Dev Group',     'Construction',  '2022-06-30', 1750000, 5.75),
    ('Riverside Medical',      'Commercial RE', '2021-08-22', 2800000, 4.90),
    ('Summit Retail Plaza',    'Commercial RE', '2020-12-01',  975000, 5.10),
    ('Valley Farm Supply',     'Agricultural',  '2023-05-17',  380000, 5.85),
    ('Westside Auto Group',    'Commercial RE', '2022-09-08', 1450000, 5.40),
    ('Northgate Condos',       'Multifamily',   '2021-05-25', 4200000, 4.65),
]
doc = Document()
doc.add_heading('Loan Terms Summary', 0)
doc.add_paragraph('As of December 31, 2024')
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
for i, h in enumerate(['Borrower','Loan Type','Origination Date','Principal ($)','Interest Rate (%)']):
    c = table.rows[0].cells[i]
    c.text = h
    c.paragraphs[0].runs[0].font.bold = True
for b, lt, od, prin, rate in borrowers:
    row = table.add_row().cells
    row[0].text = b
    row[1].text = lt
    row[2].text = od
    row[3].text = f'{prin:,.0f}'
    row[4].text = str(rate)
doc.save(f'{out}/Example2/loan_terms.docx')
print(f"Example2/loan_terms.docx: {len(borrowers)} borrowers")

pay_rows = []
for b, lt, od, prin, rate in borrowers:
    mr = rate / 100 / 12
    pmt = prin * mr / (1 - (1 + mr)**-360)
    bal = prin
    orig = date.fromisoformat(od)
    for m in range(1, 25):
        yr, mo = divmod(orig.month + m - 1, 12)
        pdate = date(orig.year + yr, mo + 1, 1)
        interest = bal * mr
        princ_paid = pmt - interest
        bal -= princ_paid
        pay_rows.append({
            'Borrower': b,
            'Payment Date': pdate.strftime('%Y-%m-%d'),
            'Payment': round(pmt, 2),
            'Principal Paid': round(princ_paid, 2),
            'Interest Paid': round(interest, 2),
            'Remaining Balance': round(max(bal, 0), 2)
        })
pd.DataFrame(pay_rows).to_excel(f'{out}/Example2/payments.xlsx', index=False)
print(f"Example2/payments.xlsx: {len(pay_rows)} rows")

# ── Example3: earnings.pdf + prices.xlsx ─────────────────────────────────────
companies = [
    ('AAPL',  'Apple Inc',          120.3, 7.26, 'Beats estimates', 222),
    ('MSFT',  'Microsoft Corp',      96.4, 3.23, 'In line',         415),
    ('GOOGL', 'Alphabet Inc',        90.2, 2.14, 'Misses estimates',191),
    ('AMZN',  'Amazon.com Inc',     159.0, 1.86, 'Beats estimates', 218),
    ('NVDA',  'NVIDIA Corp',         35.1, 0.89, 'Beats estimates', 495),
    ('META',  'Meta Platforms',      48.4, 8.02, 'Beats estimates', 589),
    ('TSLA',  'Tesla Inc',           25.7, 0.73, 'Misses estimates',248),
    ('JPM',   'JPMorgan Chase',      42.8, 4.81, 'Beats estimates', 234),
    ('JNJ',   'Johnson & Johnson',   22.5, 2.04, 'In line',         157),
    ('XOM',   'Exxon Mobil',         90.0, 2.48, 'In line',         107),
]

pdf = FPDF()
pdf.add_page()
pdf.set_font('Helvetica', 'B', 16)
pdf.cell(0, 12, 'Q4 2024 Earnings Summary', ln=True, align='C')
pdf.set_font('Helvetica', '', 11)
pdf.cell(0, 8, 'Selected S&P 500 Companies', ln=True, align='C')
pdf.ln(4)
cols   = ['Ticker', 'Company',               'Revenue ($B)', 'EPS ($)', 'Guidance']
widths = [22,        60,                       32,             22,         54]
pdf.set_font('Helvetica', 'B', 10)
pdf.set_fill_color(220, 230, 242)
for w, h in zip(widths, cols):
    pdf.cell(w, 9, h, border=1, fill=True)
pdf.ln()
pdf.set_font('Helvetica', '', 10)
for t, co, rev, eps, guid, _ in companies:
    for w, v in zip(widths, [t, co, str(rev), str(eps), guid]):
        pdf.cell(w, 8, v, border=1)
    pdf.ln()
pdf.output(f'{out}/Example3/earnings.pdf')
print(f"Example3/earnings.pdf: {len(companies)} companies")

prices = [{'Ticker': t, 'Company': co, 'Closing Price': price}
          for t, co, _, _, _, price in companies]
pd.DataFrame(prices).to_excel(f'{out}/Example3/prices.xlsx', index=False)
print(f"Example3/prices.xlsx: {len(prices)} rows")

# ── Zip ───────────────────────────────────────────────────────────────────────
zip_path = 'C:/Users/kerry/repos/mgmt675/files/module4.zip'
with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
    for ex in ['Example1', 'Example2', 'Example3']:
        for fname in os.listdir(f'{out}/{ex}'):
            zf.write(f'{out}/{ex}/{fname}', f'{ex}/{fname}')
            print(f"  zipped: {ex}/{fname}")
print(f"\nmodule4.zip created at {zip_path}")
